#!/usr/bin/env python3
"""
Form Filler - Fill original form files with user data

Features:
- Fill .docx files using python-docx
- Fill .doc files by converting to .docx first
- Convert filled documents to PDF
- Preserve original formatting
"""

import logging
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Try to import libraries
try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed, .docx filling disabled")

try:
    from pdf2docx import Converter

    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False
    logger.info("pdf2docx not installed, PDF conversion limited")


class FormFiller:
    """Fill original form files with user answers"""

    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "form_filler"
        self.temp_dir.mkdir(exist_ok=True)

    def fill_form(self, original_file: Path, answers: dict[str, Any], output_path: Path | None = None) -> Path:
        """
        Fill form with answers and return path to filled file

        Args:
            original_file: Path to original form file
            answers: Dict of {field_name: value}
            output_path: Optional output path, otherwise uses temp directory

        Returns:
            Path to filled file
        """
        if not original_file.exists():
            raise FileNotFoundError(f"Original file not found: {original_file}")

        extension = original_file.suffix.lower()

        if extension == ".docx":
            return self._fill_docx(original_file, answers, output_path)
        elif extension == ".doc":
            return self._fill_doc(original_file, answers, output_path)
        elif extension == ".pdf":
            return self._fill_pdf(original_file, answers, output_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def _fill_docx(self, docx_file: Path, answers: dict[str, Any], output_path: Path | None = None) -> Path:
        """
        Fill .docx file with answers

        Strategy:
        1. Find text patterns that look like form fields (e.g., "Họ và tên: ___", "[ ]")
        2. Replace with actual values
        3. Save filled document
        """
        if not HAS_DOCX:
            raise ImportError("python-docx required for .docx filling")

        doc = Document(str(docx_file))

        # Create a mapping of Vietnamese labels to values
        # This handles variations like "Họ và tên", "Họ tên", "Tên"
        label_to_value = {}
        for field_name, value in answers.items():
            # Try to get the label from field definitions if available
            # For now, use the field_name as-is
            label_to_value[field_name] = str(value) if value else ""

        # Also create common Vietnamese field mappings
        vietnamese_mappings = self._create_vietnamese_mappings(answers)
        label_to_value.update(vietnamese_mappings)

        # Fill paragraphs
        for paragraph in doc.paragraphs:
            self._fill_paragraph_text(paragraph, label_to_value)

        # Fill tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._fill_paragraph_text(paragraph, label_to_value)

        # Save filled document
        if output_path is None:
            output_path = self.temp_dir / f"filled_{docx_file.name}"

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        logger.info(f"Filled .docx saved to: {output_path}")
        return output_path

    def _fill_paragraph_text(self, paragraph, label_to_value: dict[str, str]):
        """Fill a paragraph with values"""
        for run in paragraph.runs:
            text = run.text
            for label, value in label_to_value.items():
                if not value:
                    continue

                # Pattern 1: "Label: ______" or "Label: ………"
                pattern1 = rf"{re.escape(label)}:\s*[_\.…\s]*"
                if re.search(pattern1, text, re.IGNORECASE):
                    text = re.sub(pattern1, f"{label}: {value}", text, flags=re.IGNORECASE)

                # Pattern 2: "Label:[ ]" or "Label: [ ]"
                pattern2 = rf"{re.escape(label)}:\s*\[\s*\]"
                if re.search(pattern2, text, re.IGNORECASE):
                    text = re.sub(pattern2, f"{label}: {value}", text, flags=re.IGNORECASE)

                # Pattern 3: Just the label followed by underscores/dots
                pattern3 = rf"\b{re.escape(label)}\s*[_\.…]{{3,}}"
                if re.search(pattern3, text, re.IGNORECASE):
                    text = re.sub(pattern3, f"{label} {value}", text, flags=re.IGNORECASE)

            run.text = text

    def _create_vietnamese_mappings(self, answers: dict[str, Any]) -> dict[str, str]:
        """
        Create Vietnamese label mappings from English field names

        Examples:
        - full_name → "Họ và tên", "Họ tên", "Tên"
        - dob → "Ngày sinh", "Sinh ngày"
        - address → "Địa chỉ", "Nơi ở"
        """
        mappings = {}

        # Common field mappings
        field_to_vietnamese = {
            "full_name": ["Họ và tên", "Họ tên", "Tên"],
            "dob": ["Ngày sinh", "Sinh ngày", "Ngày/tháng/năm sinh"],
            "birth_date": ["Ngày sinh", "Sinh ngày"],
            "id_number": ["Số CCCD", "Số CMND", "CCCD", "CMND", "Số chứng minh"],
            "address": ["Địa chỉ", "Nơi ở", "Chỗ ở"],
            "phone": ["Số điện thoại", "Điện thoại", "SĐT"],
            "email": ["Email", "E-mail", "Thư điện tử"],
            "gender": ["Giới tính", "Nam/Nữ"],
            "position": ["Chức vụ", "Vị trí"],
            "department": ["Phòng ban", "Bộ phận"],
            "company": ["Công ty", "Đơn vị"],
            "date": ["Ngày", "Ngày tháng"],
            "reason": ["Lý do", "Nguyên nhân"],
            "note": ["Ghi chú", "Chú thích"],
            "signature": ["Chữ ký", "Ký tên"],
        }

        for field_name, value in answers.items():
            if not value:
                continue

            vietnamese_labels = field_to_vietnamese.get(field_name, [])
            for vn_label in vietnamese_labels:
                mappings[vn_label] = str(value)

        return mappings

    def _fill_doc(self, doc_file: Path, answers: dict[str, Any], output_path: Path | None = None) -> Path:
        """
        Fill .doc file by converting to .docx first
        """
        # Convert .doc to .docx using LibreOffice
        with tempfile.TemporaryDirectory() as tmpdir:
            try:
                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        tmpdir,
                        str(doc_file),
                    ],
                    capture_output=True,
                    timeout=30,
                    check=False,
                )

                docx_path = Path(tmpdir) / f"{doc_file.stem}.docx"

                if result.returncode == 0 and docx_path.exists():
                    # Fill the converted .docx
                    return self._fill_docx(docx_path, answers, output_path)
                else:
                    raise RuntimeError(f"LibreOffice conversion failed: {result.stderr.decode()}")

            except Exception as e:
                logger.error(f"Failed to convert .doc to .docx: {e}")
                raise

    def _fill_pdf(self, pdf_file: Path, answers: dict[str, Any], output_path: Path | None = None) -> Path:
        """
        Fill PDF file

        Strategy:
        1. If PDF has form fields → use PyPDF2 to fill them
        2. If scanned PDF → convert to .docx, fill, convert back
        3. Fallback → overlay text on PDF
        """
        # For now, convert PDF to DOCX, fill, then convert back
        # This preserves layout better than overlaying text

        if not HAS_PDF2DOCX:
            logger.warning("pdf2docx not available, using basic PDF overlay")
            return self._fill_pdf_overlay(pdf_file, answers, output_path)

        with tempfile.TemporaryDirectory() as tmpdir:
            # Convert PDF to DOCX
            docx_path = Path(tmpdir) / f"{pdf_file.stem}.docx"
            cv = Converter(str(pdf_file))
            cv.convert(str(docx_path))
            cv.close()

            # Fill DOCX
            filled_docx = self._fill_docx(docx_path, answers)

            # Convert back to PDF using LibreOffice
            if output_path is None:
                output_path = self.temp_dir / f"filled_{pdf_file.stem}.pdf"

            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_path.parent),
                    str(filled_docx),
                ],
                check=True,
                timeout=30,
            )

            # Rename to expected output name
            converted_pdf = output_path.parent / f"{filled_docx.stem}.pdf"
            if converted_pdf.exists() and converted_pdf != output_path:
                converted_pdf.rename(output_path)

            logger.info(f"Filled PDF saved to: {output_path}")
            return output_path

    def _fill_pdf_overlay(self, pdf_file: Path, answers: dict[str, Any], output_path: Path | None = None) -> Path:
        """
        Fallback: overlay text on PDF
        This is basic and won't match original form layout perfectly
        """
        # TODO: Implement PDF text overlay using reportlab or PyPDF2
        # For now, raise NotImplementedError
        raise NotImplementedError("PDF overlay filling not yet implemented. Install pdf2docx.")

    def convert_to_pdf(self, docx_file: Path, output_path: Path | None = None) -> Path:
        """
        Convert filled .docx to PDF using LibreOffice

        Args:
            docx_file: Path to .docx file
            output_path: Optional output PDF path

        Returns:
            Path to generated PDF
        """
        if output_path is None:
            output_path = docx_file.with_suffix(".pdf")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_path.parent),
                    str(docx_file),
                ],
                check=True,
                capture_output=True,
                timeout=30,
            )

            # LibreOffice creates PDF with same base name
            generated_pdf = output_path.parent / f"{docx_file.stem}.pdf"

            if generated_pdf.exists() and generated_pdf != output_path:
                generated_pdf.rename(output_path)

            logger.info(f"Converted to PDF: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Failed to convert to PDF: {e}")
            raise


def fill_and_export(original_file_path: str, answers: dict[str, Any], output_pdf_path: str | None = None) -> Path:
    """
    Convenience function to fill form and export as PDF

    Args:
        original_file_path: Path to original form file
        answers: Dict of answers
        output_pdf_path: Optional output PDF path

    Returns:
        Path to generated PDF
    """
    filler = FormFiller()
    original_file = Path(original_file_path)

    # Fill the form
    filled_file = filler.fill_form(original_file, answers)

    # Convert to PDF if not already PDF
    if filled_file.suffix.lower() == ".pdf":
        if output_pdf_path and Path(output_pdf_path) != filled_file:
            filled_file.rename(output_pdf_path)
            return Path(output_pdf_path)
        return filled_file
    else:
        # Convert .docx to PDF
        output_path = Path(output_pdf_path) if output_pdf_path else filled_file.with_suffix(".pdf")
        return filler.convert_to_pdf(filled_file, output_path)


if __name__ == "__main__":
    # Test form filler
    import sys

    logging.basicConfig(level=logging.DEBUG)

    if len(sys.argv) < 2:
        logger.error("Usage: python form_filler.py <original_file>")
        sys.exit(1)

    test_file = Path(sys.argv[1])
    test_answers = {
        "full_name": "Nguyễn Văn A",
        "dob": "01/01/1980",
        "address": "123 Đường ABC, Quận 1, TP.HCM",
        "phone": "0901234567",
    }

    filled_pdf = fill_and_export(str(test_file), test_answers)
    logger.info(f"✓ Generated PDF: {filled_pdf}")
